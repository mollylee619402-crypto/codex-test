from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Tuple

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill

INPUT_DIR = Path("input")
OUTPUT_DIR = Path("output")
OUTPUT_FILE = OUTPUT_DIR / "project_document_extract.xlsx"

SUMMARY_LIMIT = 1000
TEXT_SCAN_LIMIT = 5000
EXCEL_PREVIEW_ROWS = 5

OUTPUT_HEADERS = [
    "序号",
    "文件名",
    "文件类型",
    "所在文件夹",
    "完整路径",
    "可能所属资料类型",
    "提取文本摘要",
    "可能的项目名称",
    "可能的建设单位",
    "可能的工程内容",
    "可能的金额/投资/概算",
    "备注",
]

IGNORED_FILENAMES = {".gitkeep", ".DS_Store", "Thumbs.db"}
SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx"}
MANUAL_CONFIRM_EXTENSIONS = {".doc", ".xls"}

FILE_TYPE_MAP = {
    ".doc": "Word 文档",
    ".docx": "Word 文档",
    ".pdf": "PDF 文档",
    ".xls": "Excel 表格",
    ".xlsx": "Excel 表格",
}

DOCUMENT_TYPE_KEYWORDS = [
    ("初步设计", "初步设计资料"),
    ("可行性研究", "可研资料"),
    ("可研", "可研资料"),
    ("批复", "批复文件"),
    ("概算", "概算资料"),
    ("实施方案", "实施方案"),
]

COLUMN_WIDTHS = {
    "A": 8,
    "B": 30,
    "C": 14,
    "D": 36,
    "E": 56,
    "F": 18,
    "G": 70,
    "H": 32,
    "I": 28,
    "J": 42,
    "K": 28,
    "L": 26,
}


@dataclass
class DocumentExtractItem:
    index: int
    file_name: str
    file_type: str
    folder: str
    full_path: str
    document_type: str
    summary: str
    project_name: str
    construction_unit: str
    project_content: str
    amount: str
    note: str = ""

    def to_row(self) -> List[object]:
        return [
            self.index,
            self.file_name,
            self.file_type,
            self.folder,
            self.full_path,
            self.document_type,
            self.summary,
            self.project_name,
            self.construction_unit,
            self.project_content,
            self.amount,
            self.note,
        ]


def ensure_directories() -> None:
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def compact_text(value: str) -> str:
    return re.sub(r"\s+", " ", normalize_text(value)).strip()


def limit_text(value: str, limit: int = SUMMARY_LIMIT) -> str:
    text = normalize_text(value)
    return text[:limit]


def should_include_file(path: Path) -> bool:
    if not path.is_file() or path.name in IGNORED_FILENAMES or path.name.startswith("~$"):
        return False
    return path.suffix.lower() in SUPPORTED_EXTENSIONS.union(MANUAL_CONFIRM_EXTENSIONS)


def get_file_type(extension: str) -> str:
    return FILE_TYPE_MAP.get(extension.lower(), "其他文件")


def guess_document_type(file_name: str) -> str:
    for keyword, document_type in DOCUMENT_TYPE_KEYWORDS:
        if keyword in file_name:
            return document_type
    return ""


def extract_docx_text(file_path: Path) -> Tuple[str, str]:
    from docx import Document

    document = Document(file_path)
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            parts.append(text)
        if len("\n".join(parts)) >= TEXT_SCAN_LIMIT:
            return "\n".join(parts), ""

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"表格{table_index}:")
        for row in table.rows:
            row_values = [normalize_text(cell.text) for cell in row.cells]
            row_text = " | ".join(value for value in row_values if value)
            if row_text:
                parts.append(row_text)
            if len("\n".join(parts)) >= TEXT_SCAN_LIMIT:
                return "\n".join(parts), ""

    return "\n".join(parts), ""


def extract_pdf_text(file_path: Path) -> Tuple[str, str]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    parts: List[str] = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if text:
            parts.append(f"第{page_index}页:\n{text}")
        if len("\n".join(parts)) >= TEXT_SCAN_LIMIT:
            break

    full_text = "\n".join(parts)
    if not compact_text(full_text):
        return "", "未提取到可复制文本，可能为扫描版 PDF，需要人工确认"
    return full_text, ""


def iter_row_values(row: Iterable[object]) -> List[str]:
    values: List[str] = []
    for cell in row:
        text = normalize_text(cell)
        if text:
            values.append(text)
    return values


def extract_xlsx_text(file_path: Path) -> Tuple[str, str]:
    wb = load_workbook(file_path, data_only=True, read_only=True)
    parts: List[str] = []

    try:
        for ws in wb.worksheets:
            parts.append(f"工作表：{ws.title}")
            rows = ws.iter_rows(values_only=True)

            header = next(rows, None)
            if header is not None:
                header_values = iter_row_values(header)
                if header_values:
                    parts.append("表头：" + " | ".join(header_values))

            preview_count = 0
            for row in rows:
                row_values = iter_row_values(row)
                if not row_values:
                    continue
                preview_count += 1
                parts.append(f"第{preview_count}行：" + " | ".join(row_values))
                if preview_count >= EXCEL_PREVIEW_ROWS:
                    break

            if len("\n".join(parts)) >= TEXT_SCAN_LIMIT:
                break
    finally:
        wb.close()

    return "\n".join(parts), ""


def extract_text(file_path: Path) -> Tuple[str, str]:
    extension = file_path.suffix.lower()
    if extension in MANUAL_CONFIRM_EXTENSIONS:
        return "", "暂不处理该旧版文件格式，需要人工确认"

    try:
        if extension == ".docx":
            return extract_docx_text(file_path)
        if extension == ".pdf":
            return extract_pdf_text(file_path)
        if extension == ".xlsx":
            return extract_xlsx_text(file_path)
    except Exception as exc:  # noqa: BLE001
        return "", f"读取失败：{exc}；需要人工确认"

    return "", "不支持的文件格式，需要人工确认"


def find_first_match(text: str, patterns: Iterable[str]) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
        if match:
            for group in match.groups():
                result = compact_text(group)
                if result:
                    return result[:120]
    return ""


def guess_project_name(text: str, file_name: str) -> str:
    from_text = find_first_match(
        text,
        [
            r"项目名称[：:\s]+([^\n。；;|]{3,120})",
            r"工程名称[：:\s]+([^\n。；;|]{3,120})",
            r"([^\n。；;]{3,80}(?:项目|工程))(?:初步设计|可行性研究|实施方案|批复|概算)",
        ],
    )
    if from_text:
        return from_text

    stem = Path(file_name).stem
    cleaned = re.sub(r"(初步设计|可行性研究|可研|实施方案|批复|概算|报告|说明|文本|资料)$", "", stem)
    return cleaned.strip(" -_—")[:120]


def guess_construction_unit(text: str) -> str:
    return find_first_match(
        text,
        [
            r"建设单位[：:\s]+([^\n。；;|]{2,80})",
            r"项目单位[：:\s]+([^\n。；;|]{2,80})",
            r"业主单位[：:\s]+([^\n。；;|]{2,80})",
            r"建设主体[：:\s]+([^\n。；;|]{2,80})",
        ],
    )


def guess_project_content(text: str) -> str:
    return find_first_match(
        text,
        [
            r"建设内容(?:及规模)?[：:\s]+([^\n|]{5,220})",
            r"工程内容[：:\s]+([^\n|]{5,220})",
            r"主要建设内容[：:\s]+([^\n|]{5,220})",
            r"项目建设内容[：:\s]+([^\n|]{5,220})",
        ],
    )


def guess_amount(text: str) -> str:
    return find_first_match(
        text,
        [
            (
                r"(?:总投资|项目总投资|工程总投资|概算总投资|投资估算|概算)"
                r"[：:\s]*([^\n。；;|]{0,40}?(?:\d[\d,，.]*\s*(?:万元|亿元|元)))"
            ),
            r"(\d[\d,，.]*\s*(?:万元|亿元|元))",
        ],
    )


def find_excel_preview_value(text: str, aliases: Iterable[str]) -> str:
    """Return the first preview-row value whose header contains one of the aliases."""
    lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    headers: List[str] = []

    for line in lines:
        if line.startswith("表头："):
            headers = [compact_text(value) for value in line.removeprefix("表头：").split("|")]
            continue
        if headers and re.match(r"第\d+行：", line):
            row_text = re.sub(r"^第\d+行：", "", line)
            values = [compact_text(value) for value in row_text.split("|")]
            for header, value in zip(headers, values):
                if value and any(alias in header for alias in aliases):
                    return value[:120]
    return ""


def build_extract_items() -> List[DocumentExtractItem]:
    files = sorted(path for path in INPUT_DIR.rglob("*") if should_include_file(path))
    items: List[DocumentExtractItem] = []

    for index, file_path in enumerate(files, start=1):
        text, note = extract_text(file_path)
        summary = limit_text(text)
        searchable_text = normalize_text(text[:TEXT_SCAN_LIMIT])
        excel_project_name = find_excel_preview_value(searchable_text, ["项目名称", "工程名称"])
        excel_construction_unit = find_excel_preview_value(
            searchable_text, ["建设单位", "项目单位", "业主单位"]
        )
        excel_project_content = find_excel_preview_value(searchable_text, ["建设内容", "工程内容"])
        excel_amount = find_excel_preview_value(searchable_text, ["总投资", "投资", "概算", "金额"])
        items.append(
            DocumentExtractItem(
                index=index,
                file_name=file_path.name,
                file_type=get_file_type(file_path.suffix),
                folder=str(file_path.parent.resolve()),
                full_path=str(file_path.resolve()),
                document_type=guess_document_type(file_path.name),
                summary=summary,
                project_name=(
                    excel_project_name
                    or (
                        guess_project_name(searchable_text, file_path.name)
                        if searchable_text
                        else ""
                    )
                ),
                construction_unit=(
                    excel_construction_unit or guess_construction_unit(searchable_text)
                ),
                project_content=excel_project_content or guess_project_content(searchable_text),
                amount=excel_amount or guess_amount(searchable_text),
                note=note,
            )
        )

    return items


def apply_basic_style(ws) -> None:
    header_font = Font(bold=True)
    header_fill = PatternFill(fill_type="solid", fgColor="D9EAF7")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    body_alignment = Alignment(vertical="top", wrap_text=True)

    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.alignment = body_alignment

    for column_letter, width in COLUMN_WIDTHS.items():
        ws.column_dimensions[column_letter].width = width

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def write_extract_excel(items: List[DocumentExtractItem]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "项目资料内容摘录"

    ws.append(OUTPUT_HEADERS)
    for item in items:
        ws.append(item.to_row())

    apply_basic_style(ws)
    wb.save(OUTPUT_FILE)


def main() -> None:
    ensure_directories()
    items = build_extract_items()
    write_extract_excel(items)

    manual_confirm_count = sum(1 for item in items if "需要人工确认" in item.note)
    print(f"扫描到可摘录/需确认文件数: {len(items)}")
    print(f"需要人工确认文件数: {manual_confirm_count}")
    print(f"输出文件路径: {OUTPUT_FILE.resolve()}")


if __name__ == "__main__":
    main()
